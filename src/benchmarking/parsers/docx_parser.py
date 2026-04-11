"""
python-docx baseline parser.

Extracts paragraphs and table cells from DOCX files.
Install: pip install python-docx
"""

from pathlib import Path
from .base_parser import BaseParser


class DocxParser(BaseParser):
    name = "python-docx"

    SUPPORTED = {".docx"}

    def __init__(self):
        import docx  # noqa: F401 — import check
        self._docx = docx

    def supports(self, path: Path) -> bool:
        return Path(path).suffix.lower() in self.SUPPORTED

    def _parse(self, path: Path) -> str:
        doc = self._docx.Document(str(path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading style as a markdown-style prefix
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name:
                    level = "".join(filter(str.isdigit, style_name)) or "1"
                    text = "#" * int(level) + " " + text
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))

        return "\n".join(parts)
